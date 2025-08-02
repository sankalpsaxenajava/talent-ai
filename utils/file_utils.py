"""File functions to read pdf, docx, download url etc."""

from io import BytesIO
import io
import os
from urllib.parse import unquote, urlparse
from docx import Document
import docx2txt
from pypdf import PdfReader
import requests
from .util import DocType
from loguru import logger
from llama_parse import LlamaParse  # pip install llama-parse
import uuid
from pypdfium2 import PdfDocument

def get_bytes_from_file(file_path: str):
    """Fetch the bytes in file."""
    in_file = open(file_path, "rb")
    filename, file_ext = os.path.splitext(file_path)
    content = in_file.read()
    bytes_data = BytesIO(content)
    logger.debug(f"{filename} [{file_ext}]")
    doc_type = DocType.INVALID
    if file_ext == ".docx":
        doc_type = DocType.DOCX
    elif file_ext == ".pdf":
        doc_type = DocType.PDF
    else:
        raise RuntimeError("Unsupported doc type for jd, ext:" + file_ext)
    return (bytes_data, doc_type)


def get_bytes_from_url(url: str):
    """
    Fetch the bytes from URL.

    NOTE: This can return an exception in case there are
    read or write failures.

    """
    logger.debug(url)
    parsed_url = urlparse(url)
    file_path = unquote(parsed_url.path)
    filename, file_ext = os.path.splitext(file_path)
    resp = requests.get(url)
    logger.debug(resp)
    bytes_data = BytesIO(resp.content)
    logger.debug(f"filename: {filename} file_ext={file_ext};")

    doc_type = DocType.INVALID
    if file_ext == ".docx":
        doc_type = DocType.DOCX
    elif file_ext == ".pdf":
        doc_type = DocType.PDF
    else:
        raise RuntimeError("Unsupported doc type for jd, ext:" + file_ext)
    return (bytes_data, doc_type)

def extract_text_from_doc2x(bytes_data):
    """Extract teh text from word (docx) doc."""

    # NOTE: This package is needed as other python-docx is not reading headers.
    plain_text = docx2txt.process(bytes_data)
    return plain_text

def extract_text_from_docx(bytes_data):
    """Extract teh text from word (docx) doc."""
    doc = Document(bytes_data)

    # Initialize a string to store the plain text content
    plain_text = ""

    # Extract text from paragraphs
    for para in doc.paragraphs:
        plain_text += para.text + "\n"

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join([cell.text.strip() for cell in row.cells])
            plain_text += row_text + "\n"
        plain_text += "\n"  # Add an extra line to separate tables

    return plain_text


def extract_text_from_pdf(bytes_data, use_pypdfium2=True):
    """Extract the text from pdf doc."""
    content = ""
    if use_pypdfium2:
        logger.info("Using pypdfium2")
        pdf_document = PdfDocument(bytes_data)
        for page_number in range(len(pdf_document)):
            page = pdf_document.get_page(page_number)
            content += page.get_textpage().get_text_range()
    else:
        pdf_reader = PdfReader(bytes_data)
        num_pages = len(pdf_reader.pages)
        for i in range(num_pages):
            page = pdf_reader.pages[i]
            content += page.extract_text()
    return content


def extract_with_llama_parse(bytes_data: BytesIO, docType: DocType, result_type="text"):
    """Extract the text using Llama Parse.

    NOTE: It supports docx and pdf so no special handling needed.
    """
    logger.debug(f"LlamaParse result_type: {result_type}")
    # "markdown" and "text" are available
    parser = LlamaParse(result_type=result_type)
    UPLOADS_FOLDER = os.getenv("WORKING_FOLDER")
    os.makedirs(UPLOADS_FOLDER, exist_ok=True)

    assert docType == DocType.PDF or docType == DocType.DOCX

    file_uuid = uuid.uuid4()
    file_ext = ".pdf" if docType == DocType.PDF else ".docx"
    unique_filename = f"{file_uuid}{file_ext}"
    logger.debug(f"filename: {unique_filename}")
    file_path = os.path.join(UPLOADS_FOLDER, unique_filename)
    content = ""
    with open(file_path, "wb") as f:
        f.write(bytes_data.getbuffer())
        logger.debug(file_path)
        
        documents = parser.load_data(file_path)
        for doc in documents:
            content = content + doc.text.strip()
    logger.debug(f"Loaded content for {file_path} from Llamaparse\n")

    try:
        os.remove(file_path)
    except Exception as ex:
        logger.warn(f"file {file_path} removal failed with {ex}")
    return content


def extract_text(bytes_data, docType: DocType, useLMParser: bool = False):
    """Extract text from document."""
    logger.debug("DocType: ", docType)

    content = ""
    if useLMParser:
        content = extract_with_llama_parse(bytes_data, docType)

    # If llamaparse fails try with regular extract.
    if len(content) <= 10:
        logger.trace("inside extract text manual ")
        if docType == DocType.PDF:
            logger.trace("inside pdf text manual ")
            content = extract_text_from_pdf(bytes_data)

        elif docType == DocType.DOCX:
            logger.trace("inside doc text manual ")
            content = extract_text_from_doc2x(bytes_data)
        else:
            logger.error("Invalid doc type")
            raise ValueError

    if len(content) <= 10 and not useLMParser:
        # Try with llamaparse if we are not using llamaparse by default
        content = extract_with_llama_parse(bytes_data, docType)    

    logger.trace("content: ", content)

    if len(content) < 10:
        raise Exception("Unable to parse document: Potential image based document.")

    return content
